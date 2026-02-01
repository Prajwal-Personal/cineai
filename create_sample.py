from docx import Document

def create_sample_script():
    doc = Document()
    
    # Scene 1
    doc.add_paragraph("EXT. ABANDONED WAREHOUSE - NIGHT")
    doc.add_paragraph("The rain beats down on the rusted roof. A lone spotlight flickers.")
    doc.add_paragraph("MARCUS (30s) stands by the edge, smoking. He looks at a BROKEN WATCH.")
    
    doc.add_paragraph("SARAH (V.O.)")
    doc.add_paragraph("(through static)")
    doc.add_paragraph("I told you we shouldn't have come here, Marcus. The perimeter is compromised.")
    
    doc.add_paragraph("MARCUS")
    doc.add_paragraph("It's too late for that, Sarah. We're already in.")
    
    doc.add_paragraph("CUT TO:")
    
    # Scene 2
    doc.add_paragraph("INT. WAREHOUSE - CONTINUOUS")
    doc.add_paragraph("Marcus steps inside. The air is thick with dust.")
    doc.add_paragraph("He pulls out a KEYCARD and swipes it. The door clicks.")
    
    doc.save("c:/Users/Prajw/SampleSmart/sample_script.docx")

if __name__ == "__main__":
    create_sample_script()
